from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document('DOCUMENTATION/TRABAJO_GRUPAL/CiberEscudo_Semana5.docx')

# Encontrar el párrafo de Referencias
ref_para = None
for p in doc.paragraphs:
    if 'Referencias' in p.text and p.style.name.startswith('Heading'):
        ref_para = p
        break


def add_before(ref_p, text, style='Normal', bold=False):
    para = doc.add_paragraph()
    try:
        para.style = doc.styles[style]
    except Exception:
        para.style = doc.styles['Normal']
    if text:
        run = para.add_run(text)
        if bold:
            run.bold = True
    para._p.getparent().remove(para._p)
    ref_p._p.addprevious(para._p)
    return para


def add_code(ref_p, text):
    para = doc.add_paragraph()
    try:
        para.style = doc.styles['No Spacing']
    except Exception:
        para.style = doc.styles['Normal']
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    para._p.getparent().remove(para._p)
    ref_p._p.addprevious(para._p)
    return para


# Salto de página
page_break_para = doc.add_paragraph()
br = OxmlElement('w:br')
br.set(qn('w:type'), 'page')
run_el = OxmlElement('w:r')
run_el.append(br)
page_break_para._p.append(run_el)
page_break_para._p.getparent().remove(page_break_para._p)
ref_para._p.addprevious(page_break_para._p)

# 5. Título principal
add_before(ref_para, '5. Guía de instalación y ejecución local', 'Heading 1')
add_before(ref_para,
    'Esta sección describe los pasos necesarios para clonar el repositorio, '
    'configurar el entorno virtual, instalar las dependencias e iniciar la '
    'aplicación CiberEscudo en un equipo local.')

# 5.1 Requisitos
add_before(ref_para, '5.1 Requisitos previos', 'Heading 2')
for item in [
    'Python 3.10 o superior — https://www.python.org/downloads/',
    'Git — https://git-scm.com/',
    'Terminal (PowerShell, CMD o bash en macOS/Linux)',
]:
    add_before(ref_para, item, 'List Bullet')

# 5.2 Clonar
add_before(ref_para, '5.2 Clonar el repositorio', 'Heading 2')
add_before(ref_para, 'Ejecuta los siguientes comandos en tu terminal:')
add_code(ref_para, 'git clone https://github.com/cr1c4rd0/POLI_ARQUITECTURA_DE_SOFTWARE.git')
add_code(ref_para, 'cd POLI_ARQUITECTURA_DE_SOFTWARE')

# 5.3 Entorno virtual
add_before(ref_para, '5.3 Crear el entorno virtual', 'Heading 2')
add_before(ref_para,
    'Se recomienda usar un entorno virtual para aislar las dependencias del proyecto '
    'y evitar conflictos con otras instalaciones de Python en el equipo.')
add_before(ref_para, 'En Windows:', bold=True)
add_code(ref_para, 'python -m venv .venv')
add_code(ref_para, '.venv\\Scripts\\activate')
add_before(ref_para, 'En macOS / Linux:', bold=True)
add_code(ref_para, 'python3 -m venv .venv')
add_code(ref_para, 'source .venv/bin/activate')
add_before(ref_para,
    'Cuando el entorno este activo, el prompt de la terminal mostrara el prefijo (.venv) al inicio de la linea.')

# 5.4 Instalar dependencias
add_before(ref_para, '5.4 Instalar Flask y dependencias', 'Heading 2')
add_before(ref_para, 'Con el entorno virtual activo, instala todos los paquetes necesarios:')
add_code(ref_para, 'pip install -r requirements.txt')
add_before(ref_para, 'El archivo requirements.txt incluye los siguientes paquetes:')
for item in [
    'flask 3.1.3 — Framework web principal.',
    'requests 2.32.3 — Consultas HTTP a la API de Have I Been Pwned.',
    'python-dotenv 1.0.1 — Carga de variables de entorno desde archivo .env.',
]:
    add_before(ref_para, item, 'List Bullet')

# 5.5 Ejecutar
add_before(ref_para, '5.5 Ejecutar el proyecto', 'Heading 2')
add_before(ref_para, 'Desde la raiz del repositorio, ejecuta:')
add_code(ref_para, 'python app.py')
add_before(ref_para, 'Deberias ver en consola una salida similar a la siguiente:')
add_code(ref_para, ' * Running on http://127.0.0.1:5000')
add_code(ref_para, ' * Debug mode: on')
add_before(ref_para,
    'Abre tu navegador y navega a http://127.0.0.1:5000 para acceder a la aplicacion.')

# 5.6 Login
add_before(ref_para, '5.6 Iniciar sesion', 'Heading 2')
add_before(ref_para,
    'CiberEscudo utiliza un usuario fijo para el entorno de desarrollo (MVP local). '
    'No es necesario registrarse. Las credenciales de acceso son:')
for item in [
    'Usuario: admin',
    'Contrasena: ciberescudo123',
]:
    add_before(ref_para, item, 'List Bullet')
add_before(ref_para,
    'Ingresa las credenciales en la pantalla de login y haz clic en Iniciar Sesion. '
    'El sistema validara las credenciales y redirigira automaticamente al panel principal de CiberEscudo.')

# 5.7 Estructura
add_before(ref_para, '5.7 Estructura del proyecto', 'Heading 2')
add_before(ref_para, 'Organizacion de carpetas y archivos principales del proyecto:')
for line in [
    'app.py                  - Punto de entrada principal de Flask.',
    'config.py               - Configuracion global y registro de servicios (SOA).',
    'requirements.txt        - Lista de dependencias Python.',
    'routes/                 - Blueprints de cada servicio SOA.',
    '    auth.py             - Servicio de autenticacion.',
    '    monitoring.py       - Servicio de monitoreo de correos.',
    '    passwords.py        - Servicio de verificacion de contrasenas.',
    '    guides.py           - Servicio de guias de accion.',
    '    history.py          - Servicio de historial de alertas.',
    'services/               - Logica de negocio de cada servicio.',
    '    hibp_service.py     - Adaptador HIBP (brechas de correo).',
    '    password_checker.py - Verificacion con k-Anonymity.',
    '    action_guides.py    - Guias en espanol con referencias colombianas.',
    '    history_service.py  - Persistencia SQLite.',
    'templates/              - Vistas HTML con Jinja2 y Bootstrap 5.',
    'static/css/             - Estilos personalizados del tema oscuro.',
]:
    add_code(ref_para, line)

# 5.8 Endpoints
add_before(ref_para, '5.8 Endpoints disponibles', 'Heading 2')
add_before(ref_para, 'Rutas HTTP expuestas por la aplicacion:')
for item in [
    'GET  /                      Panel principal (requiere sesion activa).',
    'GET/POST  /auth/login       Formulario de inicio de sesion.',
    'GET  /auth/logout           Cierre de sesion.',
    'POST /monitoring/check      Verificar si un correo fue filtrado.',
    'GET/POST  /passwords/check  Verificar si una contrasena fue comprometida.',
    'GET  /history/              Historial de alertas registradas.',
    'GET  /guides/<tipo>         Guia de accion para un tipo de dato expuesto.',
    'GET  /api/v1/health         Estado global del sistema.',
]:
    add_before(ref_para, item, 'List Bullet')

doc.save('DOCUMENTATION/TRABAJO_GRUPAL/CiberEscudo_Semana5_updated.docx')
print('Documento guardado como CiberEscudo_Semana5_updated.docx')
